"""
ResumeCraft AI Studio - Export Utilities
Export resume data to PDF, DOCX, and TXT using reportlab and python-docx.
"""

import io
import os
from typing import Dict, Optional
from datetime import datetime

# PDF
try:
    from reportlab.lib.pagesizes import A4, letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     HRFlowable, Table, TableStyle, KeepTogether)
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# DOCX
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ─────────────────────────────────────────────
# TEXT EXPORT
# ─────────────────────────────────────────────

def resume_data_to_text(resume_data: Dict) -> str:
    """Convert resume data dict to plain text format."""
    lines = []
    p = resume_data.get('personal', {})

    # Header
    name = p.get('name', '').upper()
    if name:
        lines.append(name)
        lines.append('=' * len(name))

    contact_parts = []
    if p.get('email'): contact_parts.append(p['email'])
    if p.get('phone'): contact_parts.append(p['phone'])
    if p.get('linkedin'): contact_parts.append(p['linkedin'])
    if p.get('github'): contact_parts.append(p['github'])
    if contact_parts:
        lines.append(' | '.join(contact_parts))
    if p.get('location'):
        lines.append(p['location'])
    lines.append('')

    # Summary
    if resume_data.get('summary'):
        lines.append('PROFESSIONAL SUMMARY')
        lines.append('-' * 20)
        lines.append(resume_data['summary'])
        lines.append('')

    # Education
    if resume_data.get('education'):
        lines.append('EDUCATION')
        lines.append('-' * 9)
        for edu in resume_data['education']:
            lines.append(f"{edu.get('degree', '')} | {edu.get('institution', '')} | {edu.get('year', '')}")
            if edu.get('grade'): lines.append(f"  {edu['grade']}")
            if edu.get('details'): lines.append(f"  {edu['details']}")
        lines.append('')

    # Experience
    if resume_data.get('experience'):
        lines.append('PROFESSIONAL EXPERIENCE')
        lines.append('-' * 23)
        for exp in resume_data['experience']:
            lines.append(f"{exp.get('title', '')} | {exp.get('company', '')} | {exp.get('duration', '')}")
            if exp.get('location'): lines.append(f"  {exp['location']}")
            for bullet in exp.get('bullets', []):
                lines.append(f"  \u2022 {bullet}")
        lines.append('')

    # Internships
    if resume_data.get('internships'):
        lines.append('INTERNSHIPS')
        lines.append('-' * 11)
        for exp in resume_data['internships']:
            lines.append(f"{exp.get('title', '')} | {exp.get('company', '')} | {exp.get('duration', '')}")
            for bullet in exp.get('bullets', []):
                lines.append(f"  \u2022 {bullet}")
        lines.append('')

    # Projects
    if resume_data.get('projects'):
        lines.append('PROJECTS')
        lines.append('-' * 8)
        for proj in resume_data['projects']:
            title_line = proj.get('title', '')
            if proj.get('tech'): title_line += f" | {proj['tech']}"
            if proj.get('year'): title_line += f" | {proj['year']}"
            lines.append(title_line)
            for bullet in proj.get('bullets', []):
                lines.append(f"  \u2022 {bullet}")
        lines.append('')

    # Skills
    if resume_data.get('skills'):
        lines.append('SKILLS')
        lines.append('-' * 6)
        skills = resume_data['skills']
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                if skill_list:
                    lines.append(f"  {category.title()}: {', '.join(skill_list)}")
        elif isinstance(skills, list):
            lines.append('  ' + ', '.join(skills))
        lines.append('')

    # Certifications
    if resume_data.get('certifications'):
        lines.append('CERTIFICATIONS')
        lines.append('-' * 14)
        for cert in resume_data['certifications']:
            lines.append(f"  \u2022 {cert.get('name', '')} | {cert.get('issuer', '')} | {cert.get('year', '')}")
        lines.append('')

    # Achievements
    if resume_data.get('achievements'):
        lines.append('ACHIEVEMENTS')
        lines.append('-' * 12)
        for ach in resume_data['achievements']:
            lines.append(f"  \u2022 {ach}")
        lines.append('')

    # Publications
    if resume_data.get('publications'):
        lines.append('PUBLICATIONS')
        lines.append('-' * 12)
        for pub in resume_data['publications']:
            if isinstance(pub, dict):
                lines.append(f"  \u2022 {pub.get('title', '')} \u2014 {pub.get('journal', '')} ({pub.get('year', '')})")
            else:
                lines.append(f"  \u2022 {pub}")
        lines.append('')

    # Languages
    if resume_data.get('languages'):
        lines.append('LANGUAGES')
        lines.append('-' * 9)
        langs = resume_data['languages']
        if isinstance(langs, list):
            lines.append('  ' + ', '.join(langs))
        else:
            lines.append(f"  {langs}")
        lines.append('')

    return '\n'.join(lines)


def export_txt(resume_data: Dict) -> bytes:
    """Export resume as plain text bytes."""
    text = resume_data_to_text(resume_data)
    return text.encode('utf-8')


# ─────────────────────────────────────────────
# PDF EXPORT
# ─────────────────────────────────────────────

def export_pdf(resume_data: Dict, template_name: str = 'ats_classic') -> bytes:
    """Export resume as PDF using reportlab."""
    if not REPORTLAB_AVAILABLE:
        raise ImportError('reportlab is not installed. Run: pip install reportlab')

    buffer = io.BytesIO()

    # Page setup
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=1.5*cm,
        leftMargin=1.5*cm,
        topMargin=1.5*cm,
        bottomMargin=1.5*cm
    )

    # Styles
    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        'Name',
        fontSize=18,
        fontName='Helvetica-Bold',
        alignment=TA_CENTER,
        spaceAfter=4,
        textColor=colors.HexColor('#1a1a2e')
    )
    contact_style = ParagraphStyle(
        'Contact',
        fontSize=9,
        fontName='Helvetica',
        alignment=TA_CENTER,
        spaceAfter=8,
        textColor=colors.HexColor('#555555')
    )
    section_style = ParagraphStyle(
        'Section',
        fontSize=11,
        fontName='Helvetica-Bold',
        spaceBefore=10,
        spaceAfter=3,
        textColor=colors.HexColor('#1a1a2e'),
        borderPad=0
    )
    body_style = ParagraphStyle(
        'Body',
        fontSize=9.5,
        fontName='Helvetica',
        spaceAfter=2,
        leading=14,
        textColor=colors.HexColor('#333333')
    )
    bullet_style = ParagraphStyle(
        'Bullet',
        fontSize=9.5,
        fontName='Helvetica',
        spaceAfter=2,
        leading=14,
        leftIndent=12,
        textColor=colors.HexColor('#333333'),
        bulletFontSize=9.5,
    )
    small_style = ParagraphStyle(
        'Small',
        fontSize=9,
        fontName='Helvetica-Oblique',
        spaceAfter=2,
        textColor=colors.HexColor('#666666')
    )

    story = []
    p = resume_data.get('personal', {})

    # Name
    name = p.get('name', 'Your Name')
    story.append(Paragraph(name.upper(), name_style))

    # Contact
    contact_parts = []
    if p.get('email'): contact_parts.append(p['email'])
    if p.get('phone'): contact_parts.append(p['phone'])
    if p.get('linkedin'): contact_parts.append(p['linkedin'])
    if p.get('github'): contact_parts.append(p['github'])
    if p.get('location'): contact_parts.append(p['location'])
    if contact_parts:
        story.append(Paragraph(' | '.join(contact_parts), contact_style))

    story.append(HRFlowable(width='100%', thickness=1.5, color=colors.HexColor('#1a1a2e')))

    def add_section(title):
        story.append(Paragraph(title.upper(), section_style))
        story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#cccccc')))

    def add_body(text):
        if text:
            story.append(Paragraph(text, body_style))

    def add_bullet(text):
        if text:
            story.append(Paragraph(f'\u2022 {text}', bullet_style))

    # Summary
    if resume_data.get('summary'):
        add_section('Professional Summary')
        add_body(resume_data['summary'])

    # Education
    if resume_data.get('education'):
        add_section('Education')
        for edu in resume_data['education']:
            story.append(Paragraph(
                f"<b>{edu.get('degree', '')}</b> \u2014 {edu.get('institution', '')} ({edu.get('year', '')})",
                body_style
            ))
            if edu.get('grade'): story.append(Paragraph(edu['grade'], small_style))
            if edu.get('details'): story.append(Paragraph(edu['details'], small_style))
            story.append(Spacer(1, 4))

    # Experience
    if resume_data.get('experience'):
        add_section('Professional Experience')
        for exp in resume_data['experience']:
            story.append(Paragraph(
                f"<b>{exp.get('title', '')}</b> \u2014 {exp.get('company', '')} | {exp.get('duration', '')}",
                body_style
            ))
            if exp.get('location'): story.append(Paragraph(exp['location'], small_style))
            for bullet in exp.get('bullets', []):
                add_bullet(bullet)
            story.append(Spacer(1, 6))

    # Internships
    if resume_data.get('internships'):
        add_section('Internships')
        for exp in resume_data['internships']:
            story.append(Paragraph(
                f"<b>{exp.get('title', '')}</b> \u2014 {exp.get('company', '')} | {exp.get('duration', '')}",
                body_style
            ))
            for bullet in exp.get('bullets', []):
                add_bullet(bullet)
            story.append(Spacer(1, 4))

    # Projects
    if resume_data.get('projects'):
        add_section('Projects')
        for proj in resume_data['projects']:
            title = proj.get('title', '')
            tech = proj.get('tech', '')
            year = proj.get('year', '')
            meta = ' | '.join(filter(None, [tech, year]))
            story.append(Paragraph(
                f"<b>{title}</b>" + (f" | {meta}" if meta else ""),
                body_style
            ))
            for bullet in proj.get('bullets', []):
                add_bullet(bullet)
            story.append(Spacer(1, 4))

    # Skills
    if resume_data.get('skills'):
        add_section('Skills')
        skills = resume_data['skills']
        if isinstance(skills, dict):
            for cat, skill_list in skills.items():
                if skill_list:
                    story.append(Paragraph(
                        f"<b>{cat.title()}:</b> {', '.join(skill_list)}",
                        body_style
                    ))
        elif isinstance(skills, list):
            add_body(', '.join(skills))

    # Certifications
    if resume_data.get('certifications'):
        add_section('Certifications')
        for cert in resume_data['certifications']:
            add_bullet(f"{cert.get('name', '')} | {cert.get('issuer', '')} | {cert.get('year', '')}")

    # Achievements
    if resume_data.get('achievements'):
        add_section('Achievements')
        for ach in resume_data['achievements']:
            add_bullet(str(ach))

    # Publications
    if resume_data.get('publications'):
        add_section('Publications')
        for i, pub in enumerate(resume_data['publications'], 1):
            if isinstance(pub, dict):
                add_body(f"{i}. {pub.get('title', '')} \u2014 {pub.get('journal', '')} ({pub.get('year', '')})")
            else:
                add_body(f"{i}. {pub}")

    # Languages
    if resume_data.get('languages'):
        add_section('Languages')
        langs = resume_data['languages']
        if isinstance(langs, list):
            add_body(', '.join(langs))
        else:
            add_body(str(langs))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


# ─────────────────────────────────────────────
# DOCX EXPORT
# ─────────────────────────────────────────────

def export_docx(resume_data: Dict, template_name: str = 'ats_classic') -> bytes:
    """Export resume as DOCX using python-docx."""
    if not DOCX_AVAILABLE:
        raise ImportError('python-docx is not installed. Run: pip install python-docx')

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    def set_font(run, size=10, bold=False, italic=False, color=None):
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        set_font(run, size=11, bold=True, color=(26, 26, 46))
        # Add border below
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'AAAAAA')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    def add_body_para(text, bold_part=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if bold_part:
            run = p.add_run(bold_part)
            set_font(run, size=9.5, bold=True)
            remaining = text.replace(bold_part, '', 1)
            if remaining:
                run2 = p.add_run(remaining)
                set_font(run2, size=9.5, italic=italic)
        else:
            run = p.add_run(text)
            set_font(run, size=9.5, italic=italic)
        return p

    def add_bullet_item(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        set_font(run, size=9.5)
        return p

    p = resume_data.get('personal', {})

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(3)
    name_run = name_para.add_run(p.get('name', 'Your Name').upper())
    set_font(name_run, size=18, bold=True, color=(26, 26, 46))

    # Contact line
    contact_parts = []
    for field in ['email', 'phone', 'linkedin', 'github', 'location']:
        if p.get(field): contact_parts.append(p[field])
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(6)
        contact_run = contact_para.add_run(' | '.join(contact_parts))
        set_font(contact_run, size=9, color=(80, 80, 80))

    # Summary
    if resume_data.get('summary'):
        add_section_heading('Professional Summary')
        add_body_para(resume_data['summary'])

    # Education
    if resume_data.get('education'):
        add_section_heading('Education')
        for edu in resume_data['education']:
            add_body_para(
                f" \u2014 {edu.get('institution', '')} ({edu.get('year', '')})",
                bold_part=edu.get('degree', '')
            )
            if edu.get('grade'):
                add_body_para(edu['grade'], italic=True)
            if edu.get('details'):
                add_body_para(edu['details'], italic=True)

    # Experience
    if resume_data.get('experience'):
        add_section_heading('Professional Experience')
        for exp in resume_data['experience']:
            add_body_para(
                f" \u2014 {exp.get('company', '')} | {exp.get('duration', '')}",
                bold_part=exp.get('title', '')
            )
            if exp.get('location'):
                add_body_para(exp['location'], italic=True)
            for bullet in exp.get('bullets', []):
                add_bullet_item(bullet)

    # Internships
    if resume_data.get('internships'):
        add_section_heading('Internships')
        for exp in resume_data['internships']:
            add_body_para(
                f" \u2014 {exp.get('company', '')} | {exp.get('duration', '')}",
                bold_part=exp.get('title', '')
            )
            for bullet in exp.get('bullets', []):
                add_bullet_item(bullet)

    # Projects
    if resume_data.get('projects'):
        add_section_heading('Projects')
        for proj in resume_data['projects']:
            tech = proj.get('tech', '')
            year = proj.get('year', '')
            meta = ' | '.join(filter(None, [tech, year]))
            add_body_para(
                (f' | {meta}' if meta else ''),
                bold_part=proj.get('title', '')
            )
            for bullet in proj.get('bullets', []):
                add_bullet_item(bullet)

    # Skills
    if resume_data.get('skills'):
        add_section_heading('Skills')
        skills = resume_data['skills']
        if isinstance(skills, dict):
            for cat, skill_list in skills.items():
                if skill_list:
                    p_obj = doc.add_paragraph()
                    p_obj.paragraph_format.space_after = Pt(2)
                    r1 = p_obj.add_run(f'{cat.title()}: ')
                    set_font(r1, size=9.5, bold=True)
                    r2 = p_obj.add_run(', '.join(skill_list))
                    set_font(r2, size=9.5)
        elif isinstance(skills, list):
            add_body_para(', '.join(skills))

    # Certifications
    if resume_data.get('certifications'):
        add_section_heading('Certifications')
        for cert in resume_data['certifications']:
            add_bullet_item(f"{cert.get('name', '')} | {cert.get('issuer', '')} | {cert.get('year', '')}")

    # Achievements
    if resume_data.get('achievements'):
        add_section_heading('Achievements')
        for ach in resume_data['achievements']:
            add_bullet_item(str(ach))

    # Publications
    if resume_data.get('publications'):
        add_section_heading('Publications')
        for i, pub in enumerate(resume_data['publications'], 1):
            if isinstance(pub, dict):
                add_body_para(f"{i}. {pub.get('title', '')} \u2014 {pub.get('journal', '')} ({pub.get('year', '')})")
            else:
                add_body_para(f"{i}. {pub}")

    # Languages
    if resume_data.get('languages'):
        add_section_heading('Languages')
        langs = resume_data['languages']
        add_body_para(', '.join(langs) if isinstance(langs, list) else str(langs))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def get_export_filename(resume_data: Dict, extension: str) -> str:
    """Generate a safe filename for export."""
    name = resume_data.get('personal', {}).get('name', 'Resume')
    safe_name = '_'.join(name.split()).replace(',', '')
    timestamp = datetime.now().strftime('%Y%m%d')
    return f"{safe_name}_Resume_{timestamp}.{extension}"
